from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse
import uuid, io, tempfile, os, re, requests
from docx import Document
from pdf2docx import Converter
from dotenv import load_dotenv
from pathlib import Path
from supabase import create_client, Client
from datetime import datetime
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

root = Path(__file__).resolve().parent.parent
env_path = root / ".env.local"
load_dotenv(dotenv_path=env_path)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Supabase configuration
SUPABASE_URL = os.getenv("NEXT_PUBLIC_SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")

if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
    logger.warning("Supabase credentials not found. Database operations will be disabled.")
    supabase = None
else:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    logger.info("Supabase client initialized successfully")

app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Configuration ---

MANDATORY_FRONT_MATTER = {
    "Naslovna stran na platnici": [],
    "Notranja naslovna stran v zaključnem delu": [],
    "Naslednja notranja naslovna stran": [r"naslov(na)? stran", r"univerza", r"fakulteta"],
    "Zahvala": [r"\bzahvala\b"],
    "Povzetek SI": [r"\bpovzetek\b", r"ključne besede", r"udk"],
    "Povzetek EN": [r"\babstract\b", r"keywords", r"udc"],
    "Izjava o avtorstvu": [r"izjava o avtorstvu"],
    "Kazalo vsebine": [r"kazalo vsebine"],
    "Kazalo slik": [r"kazalo slik"],
    "Kazalo grafov": [r"kazalo grafov"],
    "Kazalo tabel": [r"kazalo tabel"],
    "Seznam simbolov in kratic": [r"seznam.*simbol", r"seznam.*kratic", r"uporabljene.*kratice"],
    "Vsebina zaključnega dela": [r"\buvod\b", r"^1\."],
    "Seznam virov in literature": [r"viri in literatura", r"seznam virov"],
    "Priloge": [r"priloge"],
}

NAME_PATTERN = re.compile(r"^[A-ZČŠĐŽ][a-zčšđž]+(?:\s+[A-ZČŠĐŽ][a-zčšđž]+)+$")
TYPE_PATTERN = re.compile(r"\b(Magistrsko delo|Diplomsko delo|Doktorska disertacija|Kandidatensko delo)\b", re.IGNORECASE)
CITYDATE_PATTERN = re.compile(
    r"^[A-ZČŠĐŽ][a-zčšđžščž]+,\s*(?:januar|februar|marec|april|maj|junij|julij|avgust|september|oktober|november|december)\s+\d{4}$",
    re.IGNORECASE
)

SUBSECTION_PATTERNS = [r"cilj", r"predpostavk", r"raziskovaln", r"omejit"]
SECTION_PATTERNS = {
    "Uvod":               r"^\d+\.\s*UVOD|^UVOD",
    "Pregled literature": r"^\d+\.\s*Pregled literature|^Pregled literature",
    "Metodologija":       r"^\d+\.\s*Metodologija|^Metodologija",
    "Rezultati":          r"^\d+\.\s*Rezultati|^Rezultati",
    "Zaključek":          r"^\d+\.\s*(?:Zaključek|Sklep)|^(?:Zaključek|Sklep)",
}

ROMAN_RE = re.compile(r"^[IVXLCDM]+$", re.IGNORECASE)

# --- Supabase Helper Functions ---

async def save_document_to_supabase(filename: str, file_type: str, analysis_result: dict):
    """Save document analysis to Supabase"""
    if not supabase:
        logger.warning("Supabase not configured, skipping database save")
        return None
    
    try:
        # Extract data from analysis result
        notranja = analysis_result.get("notranja_naslovna", {})
        structure = analysis_result.get("structure_analysis", {})
        
        # Prepare document data
        document_data = {
            "filename": filename,
            "file_type": file_type,
            "title": notranja.get("title"),
            "document_type": notranja.get("type"),
            "student_name": notranja.get("student"),
            "study_program": notranja.get("program"),
            "study_direction": notranja.get("smer"),
            "mentor": notranja.get("mentor"),
            "co_mentor": notranja.get("somentor"),
            "lecturer": notranja.get("lektor"),
            "overall_score": structure.get("overall_score"),
            "total_sections": structure.get("total_sections"),
            "found_sections": structure.get("found_sections"),
            "missing_critical_count": structure.get("missing_critical"),
            "uvod_quality": structure.get("uvod_quality"),
            "front_matter_analysis": analysis_result.get("front_matter_found", {}),
            "body_sections_analysis": analysis_result.get("body_sections_found", {}),
            "missing_sections": analysis_result.get("missing_sections", []),
            "missing_body_sections": analysis_result.get("missing_body_sections", []),
            "recommendations": structure.get("recommendations", []),
            "table_of_contents": analysis_result.get("table_of_contents", []),
            "uvod_content": analysis_result.get("uvod", []),
        }
        
        # Insert document record
        document_result = supabase.table("documents").insert(document_data).execute()
        
        if document_result.data:
            document_id = document_result.data[0]["id"]
            logger.info(f"Document saved with ID: {document_id}")
            
            # Save paragraphs
            paragraphs = analysis_result.get("paragraphs", [])
            if paragraphs:
                await save_paragraphs_to_supabase(document_id, paragraphs)
            
            return document_id
        else:
            logger.error("Failed to save document to Supabase")
            return None
            
    except Exception as e:
        logger.error(f"Error saving document to Supabase: {str(e)}")
        return None

async def save_paragraphs_to_supabase(document_id: str, paragraphs: list):
    """Save document paragraphs to Supabase"""
    if not supabase:
        return
    
    try:
        # Prepare paragraph data
        paragraph_data = []
        for idx, paragraph in enumerate(paragraphs):
            paragraph_data.append({
                "document_id": document_id,
                "paragraph_order": idx,
                "paragraph_id": paragraph.get("id"),
                "paragraph_style": paragraph.get("style"),
                "content": paragraph.get("content", "")
            })
        
        # Insert paragraphs in batches (Supabase has limits)
        batch_size = 100
        for i in range(0, len(paragraph_data), batch_size):
            batch = paragraph_data[i:i + batch_size]
            result = supabase.table("document_paragraphs").insert(batch).execute()
            
            if not result.data:
                logger.error(f"Failed to save paragraph batch {i // batch_size + 1}")
        
        logger.info(f"Saved {len(paragraphs)} paragraphs for document {document_id}")
        
    except Exception as e:
        logger.error(f"Error saving paragraphs to Supabase: {str(e)}")

# --- API Endpoints ---

@app.get("/", response_class=HTMLResponse)
def index():
    return """
<html><body>
  <input type="file" id="file"
         accept="application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document"
         onchange="upload(this.files[0])">
  <pre id="output"></pre>
  <script>
    async function upload(f) {
      const out = document.getElementById('output');
      let fd = new FormData(); fd.append('file', f);
      let url = f.name.toLowerCase().endsWith('.pdf') ? '/upload-pdf' : '/upload-docx';
      try {
        let res = await fetch(url, { method: 'POST', body: fd });
        if (!res.ok) throw new Error(await res.text());
        let j  = await res.json();
        out.textContent = JSON.stringify(j, null, 2);
      } catch (e) {
        out.textContent = 'Error: ' + e;
      }
    }
  </script>
</body></html>
"""

@app.get("/health")
def health_check():
    """Health check endpoint"""
    return {"status": "ok", "message": "Service is running"}

@app.post("/upload-docx")
async def upload_docx(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(400, "Please upload a DOCX file.")
    data = await file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise HTTPException(400, f"Error reading DOCX: {e}")
    
    # Process document
    analysis_result = _process_document(doc)
    
    # Save to Supabase
    document_id = await save_document_to_supabase(file.filename, "docx", analysis_result)
    
    # Add document_id to response
    response_data = analysis_result.copy()
    response_data["document_id"] = document_id
    
    return JSONResponse(response_data)

@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Please upload a PDF file.")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(await file.read())
        pdf_path = tmp.name
    docx_path = pdf_path[:-4] + ".docx"
    try:
        conv = Converter(pdf_path)
        conv.convert(docx_path)
        conv.close()
    except Exception as e:
        os.unlink(pdf_path)
        raise HTTPException(500, f"Conversion failed: {e}")
    try:
        doc = Document(docx_path)
    except Exception as e:
        raise HTTPException(500, f"Error reading converted DOCX: {e}")
    finally:
        os.unlink(pdf_path); os.unlink(docx_path)
    
    # Process document
    analysis_result = _process_document(doc)
    
    # Save to Supabase
    document_id = await save_document_to_supabase(file.filename, "pdf", analysis_result)
    
    # Add document_id to response
    response_data = analysis_result.copy()
    response_data["document_id"] = document_id
    
    return JSONResponse(response_data)

@app.get("/documents")
async def get_documents():
    """Get list of all documents"""
    if not supabase:
        raise HTTPException(500, "Database not configured")
    
    try:
        result = supabase.table("document_summary").select("*").order("created_at", desc=True).execute()
        return {"documents": result.data}
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        raise HTTPException(500, f"Error fetching documents: {str(e)}")

@app.get("/documents/{document_id}")
async def get_document(document_id: str):
    """Get specific document with paragraphs"""
    if not supabase:
        raise HTTPException(500, "Database not configured")
    
    try:
        # Get document
        doc_result = supabase.table("documents").select("*").eq("id", document_id).execute()
        if not doc_result.data:
            raise HTTPException(404, "Document not found")
        
        # Get paragraphs
        para_result = supabase.table("document_paragraphs").select("*").eq("document_id", document_id).order("paragraph_order").execute()
        
        document = doc_result.data[0]
        document["paragraphs"] = para_result.data
        
        return {"document": document}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching document {document_id}: {str(e)}")
        raise HTTPException(500, f"Error fetching document: {str(e)}")

@app.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document and its paragraphs"""
    if not supabase:
        raise HTTPException(500, "Database not configured")
    
    try:
        result = supabase.table("documents").delete().eq("id", document_id).execute()
        if result.data:
            return {"message": "Document deleted successfully"}
        else:
            raise HTTPException(404, "Document not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {str(e)}")
        raise HTTPException(500, f"Error deleting document: {str(e)}")

# --- Core processing (unchanged) ---

def _process_document(doc: Document):
    # 1. Extract all paragraphs
    paragraphs = _extract_paragraphs(doc)

    # 2. Style special sections
    styled = _style_special_sections(paragraphs)

    # 3. Remove all TOC sections (basic TOC skip based on title)
    filtered = []
    skip = False
    for p in styled:
        txt_low = p["content"].strip().lower()
        if re.match(r"^(kazalo vsebine|kazalo slik|kazalo grafov|kazalo tabel)", txt_low):
            skip = True
            continue
        if skip and re.match(r"^\d+\.", p["content"]):
            skip = False
        if skip:
            continue
        filtered.append(p)

    # 4. Run checks on full styled content
    notranja      = _extract_notranja_info(paragraphs)
    front         = _check_front_matter(styled)
    missing_front = [s for s, ok in front.items() if not ok]
    uvod          = _extract_uvod(styled)
    body          = _check_body_sections(styled)
    missing_body  = [s for s, ok in body.items() if not ok]

    # 5. Apply TOC-based heading styles
    toc        = _extract_toc(filtered)
    final_para = _apply_toc_styles(filtered, toc)

    # 6. Remove numbered TOC entries with dot leaders and page numbers
    final_para = _filter_out_toc_entries(final_para)

    # 7. Trim everything before "Zahvala" and remove unwanted lines
    output_para = []
    saw_zahvala = False
    for p in final_para:
        text = p["content"].strip()

        # a) Wait until we hit "Zahvala"
        if not saw_zahvala:
            if re.match(r"^zahvala\b", text, re.IGNORECASE):
                saw_zahvala = True
            else:
                continue

        # b) Skip standalone Roman numerals
        if ROMAN_RE.match(text):
            continue

        # c) Skip standalone page numbers (e.g. "1", "11", "123")
        if re.fullmatch(r"\d+", text):
            continue

        output_para.append(p)

    # 8. Calculate enhanced metrics
    structure_analysis = _calculate_structure_metrics(front, body, uvod)

    return {
        "notranja_naslovna": notranja,
        "paragraphs": output_para,
        "front_matter_found": front,
        "missing_sections": missing_front,
        "uvod": uvod,
        "body_sections_found": body,
        "missing_body_sections": missing_body,
        "table_of_contents": toc,
        "structure_analysis": structure_analysis
    }


# --- Helper functions (unchanged) ---

def _calculate_structure_metrics(front_matter, body_sections, uvod):
    """Calculate comprehensive structure analysis metrics"""
    total_sections = len(front_matter) + len(body_sections)
    found_sections = sum(front_matter.values()) + sum(body_sections.values())
    
    # Base score
    base_score = (found_sections / total_sections) * 100 if total_sections > 0 else 0
    
    # Bonus points for well-structured introduction
    uvod_bonus = 0
    if uvod and len(uvod) > 0:
        uvod_text = ' '.join(uvod).lower()
        subsection_matches = sum(1 for pattern in SUBSECTION_PATTERNS 
                               if re.search(pattern, uvod_text, re.IGNORECASE))
        if subsection_matches >= 2:
            uvod_bonus = 5
        elif subsection_matches >= 1:
            uvod_bonus = 2
    
    # Critical sections penalty
    critical_sections = ["Povzetek SI", "Povzetek EN", "Uvod", "Zaključek"]
    missing_critical = sum(1 for section in critical_sections 
                          if section in {**front_matter, **body_sections} 
                          and not {**front_matter, **body_sections}[section])
    critical_penalty = missing_critical * 5
    
    final_score = max(0, min(100, base_score + uvod_bonus - critical_penalty))
    
    return {
        "overall_score": round(final_score),
        "total_sections": total_sections,
        "found_sections": found_sections,
        "missing_critical": missing_critical,
        "uvod_quality": len(uvod) if uvod else 0,
        "recommendations": _generate_recommendations(found_sections, total_sections, missing_critical, uvod)
    }

def _generate_recommendations(found, total, missing_critical, uvod):
    """Generate specific recommendations for improvement"""
    recommendations = []
    
    completion_rate = (found / total) * 100 if total > 0 else 0
    
    if completion_rate < 70:
        recommendations.append("Struktura potrebuje večje popravke - manjka več kot 30% obveznih elementov.")
    elif completion_rate < 90:
        recommendations.append("Struktura je dobra, vendar bi lahko dodali manjkajoče elemente.")
    else:
        recommendations.append("Odlična struktura dokumenta!")
    
    if missing_critical > 0:
        recommendations.append(f"Manjka {missing_critical} kritičnih sekcij (povzetek, uvod, zaključek).")
    
    if not uvod or len(uvod) < 3:
        recommendations.append("Uvod je prekratek ali manjka - priporoča se razširitev.")
    
    return recommendations

def _filter_out_toc_entries(paragraphs):
    toc_line_re = re.compile(r"^\s*\d+(\.\d+)*\.?\s+.+?\.{2,}\s*\d+\s*$")
    return [p for p in paragraphs if not toc_line_re.match(p["content"])]

def _extract_paragraphs(doc: Document):
    out = []
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            out.append({"id": str(uuid.uuid4()), "style": para.style.name, "content": txt})
    return out

def _extract_notranja_info(paragraphs):
    info = {"title": None, "type": None, "student": None,
            "program": None, "smer": None,
            "mentor": None, "somentor": None, "lektor": None}
    for idx, p in enumerate(paragraphs):
        c = p["content"].strip()
        if TYPE_PATTERN.search(c):
            info["type"] = c
            # title lines before
            title_lines = []
            j = idx - 1
            while j >= 0 and paragraphs[j]["content"].strip():
                if TYPE_PATTERN.search(paragraphs[j]["content"]):
                    break
                title_lines.insert(0, paragraphs[j]["content"].strip())
                j -= 1
            info["title"] = " ".join(title_lines)
        if c.startswith("Študent"):
            info["student"] = c.split(":",1)[1].strip()
        elif c.startswith("Študijski program"):
            prog = c.split(":",1)[1].strip()
            info["program"] = prog
            nxt = paragraphs[idx+1]["content"].strip() if idx+1<len(paragraphs) else ""
            if nxt and ":" not in nxt:
                info["program"] += " — " + nxt
        elif c.startswith("Smer"):
            info["smer"] = c.split(":",1)[1].strip()
        elif c.startswith("Mentor"):
            info["mentor"] = c.split(":",1)[1].strip()
        elif c.startswith("Somentor"):
            info["somentor"] = c.split(":",1)[1].strip()
        elif c.startswith("Lektor"):
            info["lektor"] = c.split(":",1)[1].strip()
    return info

def _style_special_sections(paragraphs):
    styled = []
    for p in paragraphs:
        t = p["content"].strip()
        low = t.lower()
        if low == "zahvala":
            p["style"] = "Heading 1"
        elif re.match(r"ključne besede[:]? ", t, re.IGNORECASE):
            p["style"] = "Heading 2"
        elif re.match(r"udk[:]? ", t, re.IGNORECASE):
            p["style"] = "Heading 2"
        elif low == "povzetek":
            p["style"] = "Heading 2"
        elif re.match(r"keywords[:]? ", t, re.IGNORECASE):
            p["style"] = "Heading 2"
        elif re.match(r"udc[:]? ", t, re.IGNORECASE):
            p["style"] = "Heading 2"
        elif low == "abstract":
            p["style"] = "Heading 2"
        styled.append(p)
    return styled

def _check_front_matter(paragraphs):
    result = {}
    for sec, pats in MANDATORY_FRONT_MATTER.items():
        if sec == "Naslovna stran na platnici":
            found = _validate_naslovna_stran(paragraphs)
        elif sec == "Notranja naslovna stran v zaključnem delu":
            found = _validate_notranja_stran(paragraphs)
        else:
            found = any(re.search(p, par["content"], re.IGNORECASE)
                        for par in paragraphs for p in pats)
        result[sec] = found
    return result

def _check_body_sections(paragraphs):
    res = {}
    uvod = _extract_uvod(paragraphs)
    uvok = bool(uvod) and all(
        any(re.search(pat, line, re.IGNORECASE) for line in uvod)
        for pat in SUBSECTION_PATTERNS
    )
    res["Uvod (s podsekcijami)"] = uvok
    for name, pat in SECTION_PATTERNS.items():
        res[name] = any(re.search(pat, par["content"], re.IGNORECASE) for par in paragraphs)
    return res

def _validate_naslovna_stran(paragraphs):
    block = [par["content"].strip() for par in paragraphs][:6]
    checks = [
        any(NAME_PATTERN.match(l) for l in block),
        any(TYPE_PATTERN.search(l) for l in block),
        any(CITYDATE_PATTERN.match(l) for l in block),
        any(len(l)>10 and not NAME_PATTERN.match(l)
            and not TYPE_PATTERN.search(l) and not CITYDATE_PATTERN.match(l)
            for l in block)
    ]
    return all(checks)

def _validate_notranja_stran(paragraphs):
    flags = [
        any(TYPE_PATTERN.search(p["content"]) for p in paragraphs),
        any("Študent(ka):" in p["content"] for p in paragraphs),
        any("Študijski program:" in p["content"] for p in paragraphs),
        any("Smer:" in p["content"] for p in paragraphs),
        any("Mentor(ica):" in p["content"] for p in paragraphs),
        any("Lektor(ica):" in p["content"] for p in paragraphs),
    ]
    return all(flags)

def _extract_toc(paragraphs):
    toc = []
    in_toc = False
    entry_re = re.compile(r"^\s*(?P<num>\d+(?:\.\d+)*)\.?\s+(?P<title>.*?)\s*\.{2,}\s*(?P<page>\d+)\s*$")
    end_re   = re.compile(r"^(kazalo slik|kazalo tabel|seznam virov|priloge)", re.IGNORECASE)
    for par in paragraphs:
        txt = par["content"]
        if not in_toc:
            if re.search(r"kazalo vsebine", txt, re.IGNORECASE):
                in_toc = True
            continue
        if end_re.match(txt):
            break
        m = entry_re.match(txt)
        if m:
            n = m.group("num")
            toc.append({"number": n, "title": m.group("title").strip(), "level": n.count(".")+1})
    return toc


def _apply_toc_styles(paragraphs, toc):
    styled = []

    # Regexes
    caption_re = re.compile(r"^(Slika|Tabela)\s*\d+:", re.IGNORECASE)
    toc_entry_re = re.compile(r"""
        ^\s*
        (?P<num>\d+(?:\.\d+)*)
        \.?\s+
        (?P<title>.*?)
        \s*\.{2,}\s*\d+\s*$
    """, re.VERBOSE)
    simple_num_re = re.compile(r"""
        ^\s*
        (?P<num>\d+(?:\.\d+)*)
        \.?
        (?=[^.\s])
        (?P<title>.+)$
    """, re.VERBOSE)
    # now match any numbered prefix with at least one dot (1.1, 5.5.4, etc.)
    subsec_re = re.compile(r"^\s*(?P<num>\d+(?:\.\d+)+)\.?\s*(?P<title>.+)$")

    # build lookup from TOC
    level_map = {(e["number"], e["title"]): e["level"] for e in toc}

    for p in paragraphs:
        txt = p["content"]

        # 1) Caption entries
        if caption_re.match(txt):
            p["style"] = "Caption"
            styled.append(p)
            continue

        # 2) Explicit numeric subsections (e.g. "1.1", "5.5.4")
        m_sub = subsec_re.match(txt)
        if m_sub:
            num = m_sub.group("num")
            level = num.count('.') + 1
            p["style"] = f"Heading {level}"
            styled.append(p)
            continue

        # 3) ToC‐style entries with dot‐leaders
        m_toc = toc_entry_re.match(txt)
        if m_toc:
            num   = m_toc.group("num")
            title = m_toc.group("title").strip()
            lvl   = level_map.get((num, title), num.count('.') + 1)
            p["style"] = f"Heading {lvl}"
            styled.append(p)
            continue

        # 4) Simple numbered headings fallback
        m_simple = simple_num_re.match(txt)
        if m_simple:
            num   = m_simple.group("num")
            title = m_simple.group("title").strip()
            if (num, title) in level_map:
                lvl = level_map[(num, title)]
                p["style"] = f"Heading {lvl}"
            elif title.isupper():
                p["style"] = f"Heading {num.count('.') + 1}"
            elif num.count('.') >= 2:
                p["style"] = f"Heading {num.count('.') + 1}"

        styled.append(p)

    return styled

def _extract_uvod(paragraphs):
    uvod = []
    in_sec = False
    for p in paragraphs:
        c = p['content']
        if re.match(r"^\d+\.\s+UVOD", c, re.IGNORECASE) or c.strip().upper() == "UVOD":
            in_sec = True
            continue
        if in_sec and re.match(r"^\d+\.\s+", c):
            break
        if in_sec:
            uvod.append(c)
    return uvod